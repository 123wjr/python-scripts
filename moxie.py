from docx import Document
import re

def is_highlighted(run):
    """
    检查一个 run 是否为高亮文本，避免因 None 值导致的异常。
    """
    try:
        return run.font.highlight_color is not None
    except ValueError:
        return False

def extract_sentence(paragraph, highlighted_run):
    """
    提取包含高亮文本的句子，使用标点符号识别句子边界。
    """
    text = paragraph.text
    run_text = highlighted_run.text
    start_idx = paragraph.text.find(run_text)  # 获取高亮文本的起始位置
    
    # 寻找句子开始位置
    sentence_start = max(
        text.rfind(punct, 0, start_idx) + 1
        for punct in ".!?;:"  # 句子结束标点
    )
    
    # 寻找句子结束位置
    sentence_end = min(
        (text.find(punct, start_idx) for punct in ".!?;:" if text.find(punct, start_idx) != -1),
        default=len(text)
    )
    
    return text[sentence_start:sentence_end].strip()

def extract_highlighted_sentences(doc):
    """
    从文档中提取所有包含高亮文本的句子，避免重复。
    """
    sentences = set()
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if is_highlighted(run):
                sentence = extract_sentence(paragraph, run)
                sentences.add(sentence)  # 使用集合避免重复
    return sentences

def save_sentences_to_docx(sentences, output_path):
    """
    将提取的句子写入新的 Word 文档。
    """
    new_doc = Document()
    for sentence in sentences:
        new_doc.add_paragraph(sentence)
    new_doc.save(output_path)

def main(input_path, output_path):
    """
    主函数：从输入文件提取高亮句子并保存到输出文件。
    """
    doc = Document(input_path)
    highlighted_sentences = extract_highlighted_sentences(doc)
    save_sentences_to_docx(highlighted_sentences, output_path)
    print(f"成功将高亮句子提取到新文档：{output_path}")

# 示例用法
input_file = "D:/OneDrive - 8yb506/文件/学科内容/语文/文言文错词.docx"
output_file = "highlighted_sentences.docx"
main(input_file, output_file)
